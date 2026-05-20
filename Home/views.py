import json
from datetime import datetime
from decimal import Decimal
import json
from decimal import Decimal
from datetime import datetime
from django.shortcuts import render
from django.http import JsonResponse
from django.utils import timezone
from django.contrib.contenttypes.models import ContentType
from django.contrib.auth.decorators import login_required
from .models import Employee, Payslip, GeneralLedger 
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import models
from django.db.models import Sum
from django.http import JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
import io
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from .models import (
    Department, Employee, Attendance, Payslip, 
    Asset, MaintenanceRecord, TransportLog, FarmingActivity, 
    Expense, GeneralLedger
)

# =========================================================
# CORE MARKETING / CORE PAGES
# =========================================================

def home(request):
    return render(request, 'index.html')

def about(request):
    return render(request, 'about.html')
def service(request):
    return render(request, 'service.html')
def contact(request):   
    return render(request, 'contact.html')

# =========================================================
# DASHBOARD
# =========================================================



def ledger_export_pdf(request):
    # Fetch your system model datasets matching template calculations
    transactions = GeneralLedger.objects.all().order_by('-timestamp')
    
    # Setup response stream buffers
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontSize=20, leading=24, textColor=colors.HexColor("#212529"))
    
    # Document Header Elements
    story.append(Paragraph("General Ledger Financial Statement", title_style))
    story.append(Spacer(1, 15))
    
    # Compile Data Matrix Table Elements
    table_data = [['Timestamp', 'Type', 'Reference', 'Department', 'Description', 'Amount']]
    
    for tx in transactions:
        tx_type = "INCOME" if tx.transaction_type == 'income' else "EXPENSE"
        amt_prefix = "+" if tx.transaction_type == 'income' else "-"
        amount_str = f"{amt_prefix} KES {tx.amount:,.2f}"
        
        table_data.append([
            tx.timestamp.strftime("%d %b %Y %H:%M"),
            tx_type,
            tx.reference_number or "---",
            tx.department.name if tx.department else "General Operations",
            tx.description or "---",
            amount_str
        ])
        
    ledger_table = Table(table_data, colWidths=[110, 60, 80, 100, 110, 90])
    
    # Apply standard tabular interface styling
    ledger_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f8f9fa")),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor("#6c757d")),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('ALIGN', (-1, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#dee2e6")),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#fdfdfd")]),
    ]))
    
    story.append(ledger_table)
    doc.build(story)
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="general_ledger_report.pdf"'
    return response
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def ledger_export_word(request):
    transactions = GeneralLedger.objects.all().order_by('-timestamp')
    
    doc = Document()
    
    # Document Structure Header Title Block
    title = doc.add_paragraph()
    title_run = title.add_run("General Ledger Master Audit Log")
    title_run.font.size = Pt(18)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph("Generated Master Cash Flow Balance Tracking Ledger Report Statement Matrix.").paragraph_format.space_after = Pt(20)
    
    # Table initialization grid setup
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Timestamp'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Reference'
    hdr_cells[3].text = 'Department'
    hdr_cells[4].text = 'Description'
    hdr_cells[5].text = 'Amount'
    
    for tx in transactions:
        row_cells = table.add_row().cells
        row_cells[0].text = tx.timestamp.strftime("%d %b %Y %H:%M")
        row_cells[1].text = tx.transaction_type.upper()
        row_cells[2].text = tx.reference_number or "---"
        row_cells[3].text = tx.department.name if tx.department else "General Operations"
        row_cells[4].text = tx.description or "---"
        
        prefix = "+" if tx.transaction_type == 'income' else "-"
        row_cells[5].text = f"{prefix} KES {tx.amount:,.2f}"
        
    # Save target contents back dynamically via content stream models
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="general_ledger_statement.docx"'
    return response
# def admin_dashboard(request):
    # Statistics for the Dashboard cards
    # context = {
        # 'total_employees': Employee.objects.filter(is_active=True).count(),
        # 'total_assets': Asset.objects.count(),
        # 'monthly_expenses': Expense.objects.filter(
            # expense_date__month=timezone.now().month
        # ).aggregate(Sum('amount'))['amount__sum'] or 0,
        
        # 'farming_income': FarmingActivity.objects.aggregate(Sum('sales_income'))['sales_income__sum'] or 0,
        
        # Recent Activities for a table
        # 'recent_ledger': GeneralLedger.objects.all().order_by('-timestamp')[:5],
        # 'active_farming': FarmingActivity.objects.filter(is_closed=False),
    # }
    # return render(request, 'admin.html', context)
from django.shortcuts import render
from django.http import JsonResponse
from django.utils import timezone
from django.db.models import Sum
from django.contrib.auth.decorators import login_required
from .models import Employee, Asset, Expense, FarmingActivity, GeneralLedger, ChatMessage

@login_required
def admin_dashboard(request):
    # Statistics for the Dashboard cards
    context = {
        'total_employees': Employee.objects.filter(is_active=True).count(),
        'total_assets': Asset.objects.count(),
        'monthly_expenses': Expense.objects.filter(
            expense_date__month=timezone.now().month
        ).aggregate(Sum('amount'))['amount__sum'] or 0,
        
        'farming_income': FarmingActivity.objects.aggregate(Sum('sales_income'))['sales_income__sum'] or 0,
        
        # Recent Activities for a table
        'recent_ledger': GeneralLedger.objects.all().order_by('-timestamp')[:5],
        'active_farming': FarmingActivity.objects.filter(is_closed=False),
        
        # Initial chat messages payload loading limit
        'chat_messages': ChatMessage.objects.all().order_by('-timestamp')[:50][::-1],
    }
    return render(request, 'admin.html', context)


@login_required
def get_chat_messages(request):
    """API endpoint running every 2-3 seconds to fetch new context updates"""
    last_id = request.GET.get('last_id', 0)
    messages = ChatMessage.objects.filter(id__gt=last_id).order_by('timestamp')
    
    message_list = [{
        'id': msg.id,
        'sender': msg.sender.username,
        'is_me': msg.sender == request.user,
        'message': msg.message,
        'time': msg.timestamp.strftime('%H:%M')
    } for msg in messages]
    
    return JsonResponse({'messages': message_list})


@login_required
def send_chat_message(request):
    """API endpoint handling AJAX message submissions"""
    if request.method == 'POST':
        msg_text = request.POST.get('message', '').strip()
        if msg_text:
            msg = ChatMessage.objects.create(sender=request.user, message=msg_text)
            return JsonResponse({
                'status': 'success',
                'message': {
                    'id': msg.id,
                    'sender': msg.sender.username,
                    'is_me': True,
                    'message': msg.message,
                    'time': msg.timestamp.strftime('%H:%M')
                }
            })
    return JsonResponse({'status': 'error'}, status=400)
# =========================================================
# DEPARTMENTS
# =========================================================
# def delete_department(request, pk):
    # deps= Department.objects.filter(id=pk)
    # deps.delete()
    # return redirect(department_list)
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages

def delete_department(request, pk):
    # 1. Fetch exactly one object or cleanly throw a 404 page if it doesn't exist
    department = get_object_or_404(Department, id=pk)
    
    # 2. Preserve name context before dropping the record for the toast notification
    dept_name = department.name 
    
    # 3. Perform standard single-instance deletion execution
    department.delete()
    
    # 4. Fire success message back to your base template toast module
    messages.success(request, f"Department '{dept_name}' has been successfully pruned from the ledger.")
    
    # 5. Redirect cleanly using the registered path name string mapping
    return redirect('department_list')
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Department

def department_list(request):
    if request.method == "POST":
        # Strip trailing/leading whitespace from user input
        name = request.POST.get('name', '').strip()
        
        if name:
            # Case-insensitive validation check to prevent duplicates
            if not Department.objects.filter(name__iexact=name).exists():
                Department.objects.create(name=name)
                messages.success(request, f"Department '{name}' added successfully!")
            else:
                messages.error(request, f"The department '{name}' already exists.")
        else:
            messages.error(request, "Department name cannot be blank.")
            
        return redirect('department_list')

    departments = Department.objects.all()
    
    return render(request, 'departments.html', {
        'departments': departments,
    })
# def department_list(request):
    # if request.method == "POST":
        # name = request.POST.get('name')
        # if name:
            # Check if it already exists to avoid unique constraint errors
            # if not Department.objects.filter(name=name).exists():
                # Department.objects.create(name=name)
                # messages.success(request, f"Department '{name}' added successfully!")
            # else:
                # messages.error(request, "This department already exists.")
        # return redirect('department_list')

    # departments = Department.objects.all()
    # choices = Department.NAME_CHOICES 
    
    # return render(request, 'departments.html', {
        # 'departments': departments,
        # 'choices': choices
    # })

# =========================================================
# EMPLOYEES & ATTENDANCE
# ========================================================
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from decimal import Decimal
from .models import Employee, Department

@login_required(login_url='/login/')
def employee_list(request):
    if request.method == "POST":
        dept_id = request.POST.get('department')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        id_number = request.POST.get('id_number')
        phone = request.POST.get('phone_number')
        salary = request.POST.get('basic_salary') or '0.00'
        hire_date = request.POST.get('hire_date')
        
        try:
            dept = Department.objects.get(id=dept_id)
            Employee.objects.create(
                department=dept,
                first_name=first_name,
                last_name=last_name,
                id_number=id_number,
                phone_number=phone,
                basic_salary=Decimal(salary),
                # parent_assignment=None, # Safety structural definition slot
                hire_date=hire_date,
                created_by=request.user
            )
            messages.success(request, f"Employee '{first_name} {last_name}' added successfully!")
        except Exception as e:
            messages.error(request, f"Error adding employee record: {e}")
        return redirect('employee_list')
        
    employees = Employee.objects.all().select_related('department').order_by('first_name')
    departments = Department.objects.all()
    return render(request, 'employees.html', {
        'employees': employees,
        'departments': departments
    })

@login_required(login_url='/login/')
def edit_employee(request, pk):
    if request.method == "POST":
        emp = get_object_or_404(Employee, id=pk)
        try:
            emp.department_id = request.POST.get('department')
            emp.first_name = request.POST.get('first_name')
            emp.last_name = request.POST.get('last_name')
            emp.id_number = request.POST.get('id_number')
            emp.phone_number = request.POST.get('phone_number')
            emp.basic_salary = Decimal(request.POST.get('basic_salary') or '0.00')
            emp.hire_date = request.POST.get('hire_date')
            emp.save()
            messages.success(request, f"Profile parameters for {emp.first_name} updated successfully.")
        except Exception as e:
            messages.error(request, f"Failed updating system entity: {e}")
            
    return redirect('employee_list')


@login_required(login_url='/login/')
def delete_employee(request, pk):
    emp = get_object_or_404(Employee, id=pk)
    full_name = f"{emp.first_name} {emp.last_name}"
    try:
        emp.delete()
        messages.success(request, f"Employee record sequence for '{full_name}' was purged from system ledgers.")
    except Exception as e:
        messages.error(request, f"Cannot prune profile instance matrix safely: {e}")
        
    return redirect('employee_list')


@login_required(login_url='/login/')
def toggle_employee_status(request, pk):
    emp = get_object_or_404(Employee, id=pk)
    emp.is_active = not emp.is_active
    emp.save()
    status_label = "activated" if emp.is_active else "deactivated"
    messages.success(request, f"Worker profile status for {emp.first_name} set to {status_label}.")
    return redirect('employee_list')
import json
import csv
from django.shortcuts import render, redirect
from django.http import JsonResponse, HttpResponse
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from .models import Employee, Attendance

@login_required(login_url='/login/')
def attendance_sheet(request):
    # Fetch date context parameters safely defaulting to present systems timeline standard
    date_str = request.GET.get('date', timezone.now().date().isoformat())
    employees = Employee.objects.filter(is_active=True).select_related('department')
    
    # Extract both status flags and optional remarks notes cleanly
    existing_attendance = Attendance.objects.filter(date=date_str)
    
    # Map data matrix cleanly to structural keys
    attendance_dict = {
        record.employee_id: {
            'present': record.present,
            'remarks': record.remarks
        } for record in existing_attendance
    }

    return render(request, 'attendance.html', {
        'employees': employees,
        'date_selected': date_str,
        'attendance_dict': attendance_dict
    })


@login_required(login_url='/login/')
def save_attendance(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            date = data.get('date')
            records = data.get('records')

            if not date or not records:
                return JsonResponse({'status': 'error', 'message': 'Missing data parameters.'}, status=400)

            for item in records:
                Attendance.objects.update_or_create(
                    employee_id=item['emp_id'],
                    date=date,
                    defaults={
                        'present': item['present'],
                        'remarks': item['remarks']
                    }
                )
            return JsonResponse({'status': 'success', 'message': 'Workforce presence patterns updated successfully.'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
            
    return JsonResponse({'status': 'error', 'message': 'Method context mapping unauthorized.'}, status=405)


@login_required(login_url='/login/')
def download_attendance(request):
    # Construct CSV output stream parameters
    date_str = request.GET.get('date', timezone.now().date().isoformat())
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="Attendance_Report_{date_str}.csv"'
    
    writer = csv.writer(response)
    # Excel Ledger Header Column Layout Schema
    writer.writerow(['Employee ID', 'First Name', 'Last Name', 'Department Code', 'Status Flag', 'Operational Remarks'])
    
    records = Attendance.objects.filter(date=date_str).select_related('employee__department')
    
    for row in records:
        status_string = "Present" if row.present else "Absent"
        writer.writerow([
            row.employee.id_number,
            row.employee.first_name,
            row.employee.last_name,
            row.employee.department.name,
            status_string,
            row.remarks or ""
        ])
        
    return response

@login_required(login_url='/login/')

def payroll_list(request):
    today = timezone.now()
    month = int(request.GET.get('month', today.month))
    year = int(request.GET.get('year', today.year))
    
    # Static list of months for human-readable display names
    MONTH_NAMES = [
        (1, 'January'), (2, 'February'), (3, 'March'), (4, 'April'),
        (5, 'May'), (6, 'June'), (7, 'July'), (8, 'August'),
        (9, 'September'), (10, 'October'), (11, 'November'), (12, 'December')
    ]
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            form_month = int(data.get('month', month))
            form_year = int(data.get('year', year))
            batch_status = data.get('status', 'pending')
            slips_payload = data.get('slips', [])
            is_generation_pass = data.get('generate_missing', False)
            
            target_date = datetime(form_year, form_month, 1).date()
            payment_date_val = timezone.now().date() if batch_status == 'paid' else None
            
            # --- ACTION A: GENERATE PAYSLIPS FOR NEW/MISSING EMPLOYEES ---
            if is_generation_pass:
                # Find active employees who don't have a payslip yet for this period
                existing_payslip_emp_ids = Payslip.objects.filter(
                    month_year__month=form_month, 
                    month_year__year=form_year
                ).values_list('employee_id', flat=True)
                
                missing_employees = Employee.objects.filter(is_active=True).exclude(id__in=existing_payslip_emp_ids)
                
                generated_count = 0
                for emp in missing_employees:
                    # Create baseline payslip record structure
                    Payslip.objects.create(
                        employee=emp,
                        month_year=target_date,
                        allowances=Decimal('0.00'),
                        deductions=Decimal('0.00'),
                        status='pending'
                    )
                    generated_count += 1
                
                return JsonResponse({
                    'status': 'success',
                    'message': f"Processed matching payroll structures. Formed {generated_count} new employee payslip entries."
                })
            
            # --- ACTION B: BULK UPDATE EXISTING ROW DATA ---
            updated_count = 0
            payslip_content_type = ContentType.objects.get_for_model(Payslip)
            
            for item in slips_payload:
                emp_id = item.get('employee_id')
                allowances = Decimal(item.get('allowances', '0.00') or '0.00')
                deductions = Decimal(item.get('deductions', '0.00') or '0.00')
                
                employee = Employee.objects.get(id=emp_id)
                
                slip, created = Payslip.objects.get_or_create(
                    employee=employee,
                    month_year=target_date,
                    defaults={
                        'allowances': allowances,
                        'deductions': deductions,
                        'status': batch_status,
                        'payment_date': payment_date_val
                    }
                )
                
                if not created:
                    slip.allowances = allowances
                    slip.deductions = deductions
                    slip.status = batch_status
                    if batch_status == 'paid' and not slip.payment_date:
                        slip.payment_date = payment_date_val
                    slip.save()
                
                # If individual payslip status transitions directly to PAID, register general ledger statement entry
                if batch_status == 'paid':
                    GeneralLedger.objects.update_or_create(
                        content_type=payslip_content_type,
                        object_id=slip.id,
                        defaults={
                            'transaction_type': 'expense',
                            'ledger_source': 'payroll',
                            'department': employee.department,
                            'amount': slip.net_pay,
                            'description': f"Payroll Disbursed: {employee} ({target_date.strftime('%B %Y')})",
                            'reference_number': f"PAY-{slip.id}",
                            'created_by': request.user
                        }
                    )
                
                updated_count += 1
                
            return JsonResponse({
                'status': 'success', 
                'message': f"Successfully saved configuration for {updated_count} individual employee payslips."
            })
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=400)

    # Filter out active records for view presentation layer
    payslips = Payslip.objects.filter(
        month_year__month=month, 
        month_year__year=year
    ).select_related('employee', 'employee__department')
    
    return render(request, 'payroll.html', {
        'payslips': payslips,
        'selected_month': month,
        'selected_year': year,
        'months': MONTH_NAMES,
        'years': range(today.year - 2, today.year + 2)
    })
def payslip_detail(request, pk):
    payslip = get_object_or_404(Payslip, pk=pk)
    return render(request, 'payslip_view.html', {'payslip': payslip})

# =========================================================
# ASSETS AND MAINTENANCE (WITH LEDGER integration)
# =========================================================
def delete_asset(request, pk):
    deli = get_object_or_404(Asset, id=pk)
    name =deli.name
    try:
        deli.delete()
        messages.success(request, f"Asset record sequence for '{name}' was purged from system ledgers.")
    except Exception as e:
        messages(request,f"Error deleting asset: {e}")
    return redirect('asset_list')
    
def asset_list(request):
    query = request.GET.get('q')
    status_filter = request.GET.get('status')
    assets = Asset.objects.all().select_related('department')
    
    if query:
        assets = assets.filter(name__icontains=query) | assets.filter(reg_number__icontains=query)
    if status_filter:
        assets = assets.filter(status=status_filter)
        
    return render(request, 'asset_list.html', {
        'assets': assets,
        'status_choices': Asset.STATUS_CHOICES
    })

@login_required
def asset_create_or_edit(request, pk=None):
    asset = get_object_or_404(Asset, pk=pk) if pk else None
    departments = Department.objects.all()

    if request.method == 'POST':
        data = request.POST
        try:
            is_new = asset is None
            if is_new:
                asset = Asset()
            
            asset.name = data.get('name')
            asset.reg_number = data.get('reg_number')
            asset.department_id = data.get('department')
            asset.purchase_value = Decimal(data.get('purchase_value') or '0.00')
            asset.purchase_date = data.get('purchase_date')
            asset.status = data.get('status')
            asset.description = data.get('description', '')
            asset.save()
            
            # If creating a brand new asset, automatically capture purchase capital expense inside general ledger
            if is_new:
                GeneralLedger.objects.create(
                    transaction_type='expense',
                    ledger_source='asset_purchase',
                    department=asset.department,
                    amount=asset.purchase_value,
                    description=f"Asset Capital Purchase: {asset.name} [{asset.reg_number}]",
                    reference_number=f"AST-PUR-{asset.id}",
                    content_object=asset,
                    created_by=request.user
                )
            
            messages.success(request, "Asset record captured and evaluated.")
            return redirect('asset_list')
        except Exception as e:
            messages.error(request, f"Error: {e}")

    return render(request, 'asset_form.html', {
        'asset': asset,
        'departments': departments,
        'status_choices': Asset.STATUS_CHOICES
    })

def maintenance_list(request):
    records = MaintenanceRecord.objects.all().select_related('asset').order_by('-service_date')
    return render(request, 'maintenance_list.html', {'records': records})

@login_required
def add_maintenance(request, asset_id=None):
    selected_asset = get_object_or_404(Asset, id=asset_id) if asset_id else None
    assets = Asset.objects.exclude(status='disposed')

    if request.method == 'POST':
        asset_id = request.POST.get('asset')
        asset = get_object_or_404(Asset, id=asset_id)
        cost_val = Decimal(request.POST.get('cost') or '0.00')
        
        # 1. Create maintenance record
        maintenance = MaintenanceRecord.objects.create(
            asset=asset,
            service_date=request.POST.get('service_date'),
            cost=cost_val,
            description=request.POST.get('description', ''),
            serviced_by=request.POST.get('serviced_by', ''),
            next_service_date=request.POST.get('next_service_date') or None
        )
        
        # 2. Automatically map maintenance expense outlay over to General Ledger
        GeneralLedger.objects.create(
            transaction_type='expense',
            ledger_source='maintenance',
            department=asset.department,
            amount=cost_val,
            description=f"Maintenance Service: {asset.name} ({maintenance.description[:40]})",
            reference_number=f"MNT-{maintenance.id}",
            content_object=maintenance,
            created_by=request.user
        )

        if asset.status == 'maintenance':
            asset.status = 'active'
            asset.save()

        messages.success(request, f"Maintenance record added for {asset.name} and pushed to Ledger.")
        return redirect('maintenance_list')

    return render(request, 'maintenance_form.html', {
        'assets': assets,
        'selected_asset': selected_asset
    })

# =========================================================
# TRANSPORT AND LOGISTICS (WITH LEDGER INTEGRATION)
# =========================================================

def transport_log_list(request):
    logs = TransportLog.objects.all().select_related('asset').order_by('-date', '-created_at')
    stats = {
        'total_fuel': logs.aggregate(Sum('fuel_cost'))['fuel_cost__sum'] or 0,
        'total_km': logs.aggregate(Sum('distance_covered'))['distance_covered__sum'] or 0,
    }
    return render(request, 'log_list.html', {'logs': logs, 'stats': stats})

@login_required
def add_transport_log(request):
    assets = Asset.objects.filter(status='active')

    if request.method == 'POST':
        asset_id = request.POST.get('asset')
        asset = get_object_or_404(Asset, id=asset_id)
        new_mileage = int(request.POST.get('mileage_reading') or 0)

        last_log = TransportLog.objects.filter(asset=asset).order_by('-mileage_reading').first()
        distance = 0
        if last_log:
            distance = max(0, new_mileage - last_log.mileage_reading)

        # 1. Create the detailed transport row item
        log = TransportLog.objects.create(
            asset=asset,
            date=request.POST.get('date'),
            fuel_cost=Decimal(request.POST.get('fuel_cost') or '0.00'),
            mileage_reading=new_mileage,
            distance_covered=distance,
            loading_fees=Decimal(request.POST.get('loading_fees') or '0.00'),
            unforeseen_expenses=Decimal(request.POST.get('unforeseen_expenses') or '0.00'),
            driver_name=request.POST.get('driver_name', ''),
            destination=request.POST.get('destination', ''),
            description=request.POST.get('description', ''),
        )

        # 2. Mirror logistics operational costs summary directly into the General Ledger
        GeneralLedger.objects.create(
            transaction_type='expense',
            ledger_source='transport',
            department=asset.department,
            amount=log.total_expenses,  # Property method summing fuel, loading & unforeseen fees
            description=f"Transport Trip Logistics Costs: Vehicle {asset.reg_number} to {log.destination}",
            reference_number=f"TRP-LOG-{log.id}",
            content_object=log,
            created_by=request.user
        )

        messages.success(request, f"Trip log for {asset.reg_number} processed and recorded to General Ledger.")
        return redirect('transport_log_list')

    return render(request, 'log_form.html', {'assets': assets})

# =========================================================
# FARMING CYCLES (CROPS & LIVESTOCK LEDGER INTEGRATION)
# =========================================================

def farming_list(request):
    activities = FarmingActivity.objects.all().order_by('-is_closed', '-date_started')
    total_revenue = activities.aggregate(Sum('sales_income'))['sales_income__sum'] or 0
    return render(request, 'activity_list.html', {'activities': activities, 'total_revenue': total_revenue})

from decimal import Decimal, InvalidOperation

@login_required
def farming_create_or_edit(request, pk=None):
    activity = get_object_or_404(FarmingActivity, pk=pk) if pk else None
    
    if request.method == 'POST':
        # 1. Separate your standard text/date fields from your financial cost fields
        text_fields = ['category', 'activity_name', 'date_started', 'closing_date', 'notes']
        cost_fields = [
            'feeds_cost', 'vet_expenses', 'land_prep_cost', 
            'labor_cost', 'fertilizer_cost', 'seed_cost', 'sales_income'
        ]
        
        if not activity:
            activity = FarmingActivity()
            
        # Process standard text/date fields
        for field in text_fields:
            val = request.POST.get(field)
            if val == "" and field in ['closing_date', 'notes']:
                val = None
            setattr(activity, field, val)
            
        # 2. Process financial fields and explicitly cast them to Decimal figures
        for field in cost_fields:
            val = request.POST.get(field, '0.00').strip()
            if val == "":
                val = '0.00'
            try:
                setattr(activity, field, Decimal(val))
            except (InvalidOperation, ValueError):
                setattr(activity, field, Decimal('0.00'))
            
        activity.is_closed = 'is_closed' in request.POST
        activity.save()  # Now activity.total_costs handles pure numeric math safely!

        # Resolve ledger routing classifications (Crops vs Livestock)
        ledger_source_tag = 'farming_crops' if activity.category == 'crops' else 'farming_livestock'
        farming_department = Department.objects.filter(name__icontains='Crops').first()
        # Clean old ledger logs for this project using explicit tracking reference IDs
        GeneralLedger.objects.filter(
            reference_number__in=[f"FRM-EXP-{activity.id}", f"FRM-INC-{activity.id}"]
        ).delete()

        # 3. Safe numeric comparisons
        if activity.total_costs > 0:
            GeneralLedger.objects.create(
                transaction_type='expense',
                ledger_source=ledger_source_tag,
                department=farming_department,
                amount=activity.total_costs,
                description=f"Operational Costs: Production Cycle '{activity.activity_name}'",
                reference_number=f"FRM-EXP-{activity.id}",
                created_by=request.user
            )

        if activity.sales_income > 0:
            GeneralLedger.objects.create(
                transaction_type='income',
                ledger_source=ledger_source_tag,
                department=farming_department,
                amount=activity.sales_income,
                description=f"Market Revenue Settlement: Cycle '{activity.activity_name}'",
                reference_number=f"FRM-INC-{activity.id}",
                created_by=request.user
            )
        
        messages.success(request, f"Farming activity details for '{activity.activity_name}' synchronized across ledger lines.")
        return redirect('farming_list')

    return render(request, 'activity_form.html', {'activity': activity})
# =========================================================
# STANDARD MANUAL EXPENSES & BALANCES GENERAL LEDGER
# =========================================================

def expense_list(request):
    expenses = Expense.objects.all().order_by('-expense_date')
    departments = Department.objects.all()
    
    dept_id = request.GET.get('department')
    if dept_id:
        expenses = expenses.filter(department_id=dept_id)

    total_spent = expenses.aggregate(Sum('amount'))['amount__sum'] or 0
    return render(request, 'expense_list.html', {
        'expenses': expenses,
        'departments': departments,
        'total_spent': total_spent
    })

@login_required
def expense_create(request):
    departments = Department.objects.all()
    
    if request.method == 'POST':
        try:
            expense = Expense.objects.create(
                department_id=request.POST.get('department'),
                category=request.POST.get('category'),
                amount=Decimal(request.POST.get('amount') or '0.00'),
                expense_date=request.POST.get('expense_date'),
                description=request.POST.get('description', ''),
                recorded_by=request.user  
            )
            
            GeneralLedger.objects.create(
                transaction_type='expense',
                ledger_source='general_expense',
                department_id=request.POST.get('department') or None,
                amount=expense.amount,
                description=f"Expense ({expense.category}): {expense.description}",
                reference_number=f"EXP-{expense.id}", 
                content_object=expense,
                created_by=request.user
            )
            
            messages.success(request, "Expense recorded and pushed to General Ledger successfully.")
            return redirect('expense_list')
        except Exception as e:
            messages.error(request, f"Error: {e}")

    return render(request, 'expense_form.html', {'departments': departments})

def ledger_list(request):
    transactions = GeneralLedger.objects.all().order_by('-timestamp')
    
    total_income = transactions.filter(transaction_type='income').aggregate(Sum('amount'))['amount__sum'] or 0
    total_expense = transactions.filter(transaction_type='expense').aggregate(Sum('amount'))['amount__sum'] or 0
    net_balance = total_income - total_expense

    return render(request, 'ledger_list.html', {
        'transactions': transactions,
        'total_income': total_income,
        'total_expense': total_expense,
        'net_balance': net_balance,
    })

@login_required
def ledger_create(request):
    departments = Department.objects.all()
    if request.method == 'POST':
        GeneralLedger.objects.create(
            transaction_type=request.POST.get('transaction_type'),
            ledger_source=request.POST.get('ledger_source', 'general_expense'),
            department_id=request.POST.get('department') or None,
            amount=Decimal(request.POST.get('amount') or '0.00'),
            description=request.POST.get('description'),
            reference_number=request.POST.get('reference_number'),
            created_by=request.user
        )
        return redirect('ledger_list')
    
    return render(request, 'ledger_form.html', {
        'departments': departments,
        'ledger_source_choices': GeneralLedger.ledger_source
    })
import csv
from decimal import Decimal
from django.shortcuts import render
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from django.db.models import Sum
from django.utils import timezone
from .models import GeneralLedger, Department

@login_required
def financial_statements_dashboard(request):
    # 1. Handle Optional Date Filtering
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    dept_id = request.GET.get('department')

    ledger_qs = GeneralLedger.objects.all()
    if start_date:
        ledger_qs = ledger_qs.filter(timestamp__date__gte=start_date)
    if end_date:
        ledger_qs = ledger_qs.filter(timestamp__date__lte=end_date)
    if dept_id:
        ledger_qs = ledger_qs.filter(department_id=dept_id)

    # 2. Compute Cash Flow Aggregations
    total_income = ledger_qs.filter(transaction_type='income').aggregate(Sum('amount'))['amount__sum'] or Decimal('0.00')
    total_expense = ledger_qs.filter(transaction_type='expense').aggregate(Sum('amount'))['amount__sum'] or Decimal('0.00')
    net_cash_flow = total_income - total_expense

    # 3. Compute Balance Sheet Structural Snapshot
    # Historical retained metrics govern equity position balances
    retained_earnings = total_income - total_expense
    
    cash_asset = retained_earnings if retained_earnings > 0 else Decimal('0.00')
    bank_overdraft = abs(retained_earnings) if retained_earnings < 0 else Decimal('0.00')
    
    total_assets = cash_asset
    total_equity_liabilities = retained_earnings + bank_overdraft if retained_earnings < 0 else retained_earnings

    # 4. HANDLE EXPORT/DOWNLOAD REQUEST
    if 'download' in request.GET:
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="Financial_Statement_{timezone.now().strftime("%Y%m%d")}.csv"'
        
        writer = csv.writer(response)
        writer.writerow(['FINANCIAL STATEMENT EXPORT', timezone.now().strftime('%d %B %Y')])
        writer.writerow([])
        
        # Section A: Cash Flow
        writer.writerow(['1. CASH FLOW STATEMENT'])
        writer.writerow(['Line Item', 'Amount (KES)'])
        writer.writerow(['Total Cash Inflows (Income)', total_income])
        writer.writerow(['Total Cash Outflows (Expenses)', total_expense])
        writer.writerow(['NET CASH FLOW POSITION', net_cash_flow])
        writer.writerow([])
        
        # Section B: Balance Sheet
        writer.writerow(['2. BALANCE SHEET SNAPSHOT'])
        writer.writerow(['Category', 'Line Item', 'Amount (KES)'])
        writer.writerow(['Assets', 'Current Cash/Bank Balance', cash_asset])
        writer.writerow(['Assets', 'TOTAL ENTERPRISE ASSETS', total_assets])
        writer.writerow(['Equity & Liab', 'Retained Earnings (Equity)', retained_earnings])
        writer.writerow(['Equity & Liab', 'Bank Overdraft (Short-term Liability)', bank_overdraft])
        writer.writerow(['Equity & Liab', 'TOTAL EQUITY AND LIABILITIES', total_equity_liabilities])
        
        return response

    # Context items for UI rendering
    departments = Department.objects.all()
    
    return render(request, 'financial_statements.html', {
        'total_income': total_income,
        'total_expense': total_expense,
        'net_cash_flow': net_cash_flow,
        'cash_asset': cash_asset,
        'bank_overdraft': bank_overdraft,
        'total_assets': total_assets,
        'retained_earnings': retained_earnings,
        'total_equity_liabilities': total_equity_liabilities,
        'departments': departments,
        'start_date': start_date,
        'end_date': end_date,
        'selected_dept': dept_id
    })